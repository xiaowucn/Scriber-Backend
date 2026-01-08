from io import BytesIO

import docx
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor

from remarkable.plugins.cgs.common.utils import split_suggestion


def export_docx(head_results):
    document = docx.Document()
    document.styles["Normal"].font.name = "宋体"
    document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for idx, (head, results) in enumerate(head_results):
        if idx != 0:
            document.add_page_break()
        document.add_heading(head)
        for result in results:
            if result.suggestion:
                name_paragraph = document.add_paragraph()
                name_paragraph.paragraph_format.first_line_indent = Cm(0.74)
                run = name_paragraph.add_run(f"{result.name}")
                run.font.color.rgb = RGBColor(154, 189, 230)

                for item in split_suggestion(result.suggestion):
                    paragraph = document.add_paragraph(item)
                    paragraph.paragraph_format.first_line_indent = Cm(0.74)
                document.add_paragraph("\n")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
