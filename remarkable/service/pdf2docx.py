import base64
import copy
import itertools
import logging
import os
import re
from collections import OrderedDict, defaultdict
from enum import IntEnum, unique
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import (
    WD_ALIGN_VERTICAL,
    WD_ROW_HEIGHT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_PARAGRAPH_ALIGNMENT,
)
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from interval import Interval
from mako.template import Template

from remarkable.config import project_root
from remarkable.pdfinsight.reader import PdfinsightReader
from remarkable.service.comment import (
    add_comments_in_elements,
    get_comment_position_from_interdoc,
    render_comment,
    replace_comments,
)

logger = logging.getLogger(__name__)

warning_element = [
    {
        "orient": 0,
        "font-size": "50rem",
        "text-align": "center",
        "text-indent": "2px",
        "margin-left": "0px",
        "anchor_id": "0-1",
        "comment_start": None,
        "comment_end": None,
        "margin-top": "30px",
        "margin-bottom": "30px",
        "index": 1,
    },
    "此处有无法导出到word的内容",
    [
        {
            "style": {
                "fontcolor": 15417396,
            },
            "text": "此处有无法导出到word的内容",
        }
    ],
    [],
    "paragraphs",
]


@unique
class PageOrientation(IntEnum):
    PORTRAIT = 0  # 纵向
    LANDSCAPE = 1  # 横向


def replace_space_in_cells(cell_text):
    chars = cell_text.strip()
    res = []
    for idx, char in enumerate(chars):
        if char != " ":
            res.append(char)
        else:
            if is_chinese_char_or_digit(chars[idx - 1]) and is_chinese_char_or_digit(chars[idx + 1]):
                continue
            res.append(" ")

    return "".join(res)


def is_chinese_char_or_digit(char):
    if char.isdigit() or char in [".", ",", "，"]:
        return True
    return bool(char == " " or "\u4e00" <= char <= "\u9fff")


class ExportHTML:
    THRESHOLD_DELTA = 20
    P_NUM = re.compile(r"^\d{8,}$")
    P_WHITESPACE_PATTERN = re.compile(r"[\s　]+", re.S)
    S_PUNC = "。？！，、；：“”‘’「」『』（）[]〔〕【】—…—-‐～·《》〈〉﹏_<>[](){}?!'\",;:."
    P_CATALOG = re.compile(r"(?P<name>(?:.{1,5}?)?(?:[\d.:]+)?\s*.+?)\.{10,}\s*(?P<page>\d+)$")
    P_SENTENCE_SPLIT = re.compile(r"(?:. )|。")

    @staticmethod
    def construct_elements_before_plot(target_elements, all_file=False, translate=False, config=None):
        all_block_list = []
        sent_anchor_id_dup_checker = set()
        for idx, block in enumerate(target_elements):
            if not block:
                continue
            style_body = {"orient": block.get("orient", 0)}
            pre_block = target_elements[idx - 1]
            block = ExportHTML.set_space_before(idx, block, pre_block)

            if block["element_type"] == "tables":
                table_list = ExportHTML.convert_table2cells_list(block, config)
                style_body["margin-top"] = f"{int(block.get('space_before', 3))}px"
                style_body["anchor_id"] = f"{block.get('anchor_id', '')}"
                style_body["index"] = block.get("index")
                style_body["grid"] = block.get("grid", {})
                style_body["merged_headers"] = block.get("merged_headers", {})
                style_body["comment_start"] = block.get("comment_start")
                style_body["comment_end"] = block.get("comment_end")
                all_block_list.append([style_body, table_list, "table"])
            elif block["element_type"] in ("images", "shapes", "equations", "algorithms"):
                style_body["margin-top"] = f"{int(block.get('space_before', 3))}px"
                style_body["anchor_id"] = f"{block.get('anchor_id', '')}"
                style_body["index"] = block.get("index")
                style_body["width"] = Inches((block["outline"][2] - block["outline"][0]) / 72)
                style_body["height"] = Inches((block["outline"][3] - block["outline"][1]) / 72)
                if all_file:
                    all_block_list.append([style_body, block.get("data"), "image"])
                else:
                    all_block_list.append([style_body, str(block.get("data") or b"", encoding="utf8"), "image"])
            elif block["element_type"] in [
                "paragraphs",
                "page_headers",
                "page_footers",
                "bibitems",
                "captions",
                "titles",
                "authors",
                "pairs",
                "infographic",
            ]:
                if "chars" in block and block.get("chars"):
                    span_list = ExportHTML.convert_para2span_list(block)
                else:
                    continue
                if all_file:
                    style_body["font-size"] = f"{int(block.get('styles', {}).get('fontsize', 12)) / 10}rem"
                else:
                    style_body["font-size"] = f"{int(block.get('styles', {}).get('fontsize', 12))}px"
                style_body["text-align"] = block.get("styles", {}).get("align", "left")
                style_body["text-indent"] = f"{int(block.get('styles', {}).get('text_indent', 2))}px"
                style_body["margin-left"] = f"{int(block.get('styles', {}).get('padding_left', 0))}px"
                style_body["anchor_id"] = f"{block.get('anchor_id', '')}"
                style_body["outline"] = block["outline"]
                style_body["comment_start"] = block.get("comment_start")
                style_body["comment_end"] = block.get("comment_end")
                style_body["margin-top"] = f"{int(block.get('space_before', 5))}px"
                style_body["index"] = block.get("index")
                if block.get("styles", {}).get("level"):
                    style_body["level"] = block["styles"]["level"]
                if ExportHTML.P_CATALOG.match(block.get("text", "")):
                    style_body["style_class"] = "article-contents"
                if block.get("attribute"):
                    style_body["attribute"] = block.get("attribute")
                if translate:  # should be not translate?
                    all_block_list.append([style_body, block.get("text", ""), span_list, block["element_type"]])
                else:
                    for sent in block.get("translated", []):
                        sent_anchor_id = base = (
                            f"""{block["page"]}-{"-".join([str(int(each)) for each in sent["chars"][0]["box"]])}"""
                        )
                        offset = 1
                        while sent_anchor_id in sent_anchor_id_dup_checker:
                            sent_anchor_id = f"{base}-{offset}"
                            offset += 1
                        sent_anchor_id_dup_checker.add(sent_anchor_id)
                        sent["anchor_id"] = sent_anchor_id
                    all_block_list.append(
                        [
                            style_body,
                            block.get("text", ""),
                            span_list,
                            block.get("translated", []),
                            block["element_type"],
                        ]
                    )

        return all_block_list

    @staticmethod
    def set_space_before(idx, block, pre_block):
        if idx == 0:
            return block

        if block["page"] == pre_block["page"]:
            pre_bottom = pre_block["outline"][-1]
            top = block["outline"][1]
            if top > pre_bottom:  # 左右并列的块可能top < pre_bottom
                block["space_before"] = top - pre_bottom

        # 每页和每一个分栏的第一个元素块沿用前一个元素块的margin-top
        if "space_before" not in block and pre_block.get("space_before"):
            if block.get("column_top") and block["element_type"] not in ["page_headers", "page_footers"]:
                block["space_before"] = pre_block["space_before"]

        return block

    @staticmethod
    def fix_signed_para_at_end(all_block_list, columns):
        if not columns:
            return all_block_list

        paragraph_index = []
        for idx, block in enumerate(all_block_list):
            if block[-1] in ["paragraphs", "page_headers", "page_footers"]:
                paragraph_index.append(idx)

        for idx in paragraph_index[:-3:-1]:
            block = all_block_list[idx]
            para_page = block[0]["anchor_id"].split("-")[0]
            page_width = columns[para_page]["outline"][2]
            para_style = block[0]
            if page_width - para_style["outline"][2] < para_style["outline"][0]:  # 段落距离页面右边更近
                para_style["margin-right"] = "2em"
                para_style["text-align"] = "right"
                para_style.pop("margin-left")

        return all_block_list

    @staticmethod
    def adjust_char_color(all_block_list):
        for block in all_block_list:
            ele_type = block[-1]
            if ele_type in ["paragraphs", "page_headers", "page_footers"]:
                for text in block[2]:
                    style = text.get("style")
                    ExportHTML.adjust_color(style)
            if ele_type == "table":
                for row in block[1]:
                    for cell in row:
                        style = cell.get("styles")
                        ExportHTML.adjust_table_char_color(style)

        return all_block_list

    @staticmethod
    def adjust_color(style):
        if not style:
            return
        font_color = style.get("fontcolor")
        if not font_color:
            return
        is_font_white = ExportHTML.is_color_close_to_white(font_color)
        if is_font_white:
            bg_color = style.get("bg_color")
            if bg_color:
                is_bg_white = ExportHTML.is_color_close_to_white(bg_color)
                style["fontcolor"] = bg_color if not is_bg_white else 0
            else:
                style["fontcolor"] = 0

    @classmethod
    def adjust_table_char_color(cls, style):
        if not style:
            return
        font_color = style.get("fontcolor")
        if not font_color:
            return
        # bg_color = style.get('bg_color')
        # if bg_color:
        #     font_rgb = cls.get_rgb(font_color)
        #     bg_rgb = cls.get_rgb(bg_color)
        #     delta_e = cls.color_difference(font_rgb, bg_rgb)
        #     if delta_e < cls.THRESHOLD_DELTA:  # 表格中字体与背景色相近时,将字体设置为与背景色反差最大的颜色
        #         style['fontcolor'] = cls.most_different_color_from_basic(bg_color)
        # else:
        cls.adjust_color(style)

    @classmethod
    def is_color_close_to_white(cls, color):
        red, green, blue = cls.get_rgb(color)
        yuv = red * 0.299 + green * 0.587 + blue * 0.114

        return yuv > 192

    @staticmethod
    def get_rgb(color):
        if isinstance(color, int):
            rgb = hex(color).replace("0x", "")
            rgb = "0" * (6 - len(rgb)) + rgb
        elif isinstance(color, str) and color.startswith("#"):
            rgb = color.replace("#", "")
        red = int(rgb[0:2], 16)
        green = int(rgb[2:4], 16)
        blue = int(rgb[4:6], 16)
        return red, green, blue

    @staticmethod
    def color_difference(color1, color2):
        # FIXME: colormath已经不再维护,后续需要替换
        pass
        # color1 = sRGBColor(color1[0], color1[1], color1[2], is_upscaled=True)
        # color2 = sRGBColor(color2[0], color2[1], color2[2], is_upscaled=True)
        # color1_lab = convert_color(color1, LabColor)
        # color2_lab = convert_color(color2, LabColor)
        # delta_e = delta_e_cie2000(color1_lab, color2_lab)
        # return delta_e

    @staticmethod
    def convert_para2span_list(paragraph):
        span_list = []
        chars = paragraph.get("chars")
        text = paragraph.get("text") if paragraph.get("text") else paragraph.get("value", "")
        start_index = 0
        if paragraph.get("modified") is True:
            for each in chars:
                span_list.append({"style": each, "text": each["text"]})
        else:
            _match_result = ExportHTML.P_CATALOG.match(text)
            if _match_result:
                _text_list = [_match_result["name"], _match_result["page"]]
                for each in _text_list:
                    char_idx = text.index(each[0])
                    span_list.append({"style": chars[char_idx], "text": each})
                span_list.insert(1, {"style": {}, "text": ""})
            else:
                for index, char in enumerate(chars):
                    if char["text"] in ExportHTML.S_PUNC:
                        continue
                    if ExportHTML.is_same_style(chars[start_index], char):
                        continue
                    span_list.append({"style": chars[start_index], "text": text[start_index:index]})
                    start_index = index
                if chars:
                    span_list.append({"style": chars[start_index], "text": text[start_index : len(chars)]})

        return span_list

    @staticmethod
    def is_same_style(char1, char2):
        return char1.get("style", "") == char2.get("style", "")

    @staticmethod
    def convert_table2cells_list(table_block, config):
        table_list = ExportHTML.get_table_list(table_block, config)
        for merge_data in table_block["merged"]:
            target_cell = merge_data[0]

            rowspan = set()
            colspan = set()
            for each_merge in merge_data:
                rowspan.add(each_merge[0])
                colspan.add(each_merge[1])
            try:
                table_list[target_cell[0]][target_cell[1]]["rowspan"] = len(rowspan)
                table_list[target_cell[0]][target_cell[1]]["colspan"] = len(colspan)
            except IndexError:
                continue

        return table_list

    @staticmethod
    def remove_useless_field(each_cells):
        if "chars" in each_cells:
            each_cells.pop("chars")
        if "box" in each_cells:
            each_cells.pop("box")
        if "page" in each_cells:
            each_cells.pop("page")
        if "styles_diff" in each_cells:
            each_cells.pop("styles_diff")

    @staticmethod
    def check_remove_break_space(config):
        rm_line_break = False
        rm_space = False
        if config:
            if config["line_break"] or config["line_break"] == "true":
                rm_line_break = True
            if config["space"] or config["space"] == "true":
                rm_space = True

        return rm_line_break, rm_space

    @staticmethod
    def get_table_list(table_block, config):
        rm_line_break, rm_space = ExportHTML.check_remove_break_space(config)
        row_count = len(table_block["grid"]["rows"])
        column_count = len(table_block["grid"]["columns"])
        table_list = []
        for each_row in range(row_count + 1):
            row_list = []
            for each_column in range(column_count + 1):
                if f"{str(each_row)}_{str(each_column)}" in table_block["cells"]:
                    each_cells = table_block["cells"][f"{str(each_row)}_{str(each_column)}"]
                    ExportHTML.remove_useless_field(each_cells)
                    if each_cells["styles"].get("italic"):
                        each_cells["styles"]["italic"] = "italic"
                    else:
                        each_cells["styles"]["italic"] = "normal"
                    text = each_cells.get("text") if each_cells.get("text") else each_cells.get("value", "")
                    if len(text) > 20:
                        each_cells["styles"]["align"] = "left"
                    each_cells["is_num"] = bool(ExportHTML.P_NUM.search(ExportHTML.P_WHITESPACE_PATTERN.sub("", text)))
                    if each_cells["is_num"]:
                        text = ExportHTML.P_WHITESPACE_PATTERN.sub("", text)
                    else:
                        text = text.replace("\n", "") if rm_line_break else text
                        if " " in text and rm_space:
                            text = replace_space_in_cells(text)
                    if "text" in each_cells:
                        each_cells["text"] = text
                    else:
                        each_cells["value"] = text
                    row_list.append(table_block["cells"][f"{str(each_row)}_{str(each_column)}"])
                else:
                    row_list.append({})
            table_list.append(row_list)

        return table_list

    @staticmethod
    def get_slice_index(text_list, retry=False):
        slice_index_list = []
        tmp_chars_counts = 0
        _boundary = 400 if retry else 800
        for idx, each in enumerate(text_list):
            if tmp_chars_counts > _boundary:
                if not slice_index_list:
                    slice_index_list.append((0, idx - 1))
                else:
                    slice_index_list.append((slice_index_list[-1][1], idx - 1))
                tmp_chars_counts = 0
            tmp_chars_counts += len(each)
        if slice_index_list:
            if slice_index_list[-1][1] != len(text_list):
                slice_index_list.append((slice_index_list[-1][1], len(text_list) + 1))
        else:
            slice_index_list.append((0, len(text_list) + 1))

        return slice_index_list

    @staticmethod
    def make_translate(input_text, trans_map):
        input_text = input_text.strip()
        if "\n" in input_text:
            _tmp_text_list = []
            for each in input_text.split("\n"):
                _tmp_text_list.append(trans_map.get(each, ""))
            return "\n".join(_tmp_text_list)
        return trans_map.get(input_text, "")

    @staticmethod
    def fill_translated_text(trans_map, sorted_all_elements):
        sentence_chars_info = []
        for each in sorted_all_elements:
            if each["element_type"] == "paragraphs":
                _full_text_list, _ = ExportHTML.split_chars_on_sentence(each.get("chars", []), page=each["page"])
                translated = []
                for sen in _full_text_list:
                    text = "".join(char["text"] for char in sen)
                    translated.append(
                        {
                            "sentence": text,
                            "trans_sen": ExportHTML.make_translate(text, trans_map),
                            "chars": sen,
                        }
                    )
                    if text:
                        sentence_chars_info.append(
                            {
                                "sentence": text,
                                "box": [{"page": char["page"], "box": char["box"]} for char in sen],
                                "sentence_anchor": f"""
                                {each["page"]}-{"-".join([str(int(each)) for each in sen[0]["box"]])}
                                """,
                            }
                        )
                each["translated"] = translated
            elif each["element_type"] == "tables":
                each["translate_title"] = ExportHTML.make_translate(each.get("title", ""), trans_map)
                each["translate_unit"] = ExportHTML.make_translate(each.get("unit", ""), trans_map)
                for _cells in each.get("cells", {}).values():
                    _cells["translate_text"] = ExportHTML.make_translate(_cells.get("text", ""), trans_map)

        return sentence_chars_info

    @staticmethod
    def translate_syllabuses_text(trans_map, syllabuses_list):
        for _sylla in syllabuses_list:
            _sylla["translate_text"] = ExportHTML.make_translate(_sylla.get("title", ""), trans_map)

    @staticmethod
    def split_chars_on_sentence(chars, page=None):
        full_text_list = []
        total_chars_count = 0
        collected = []
        for char, next_char in itertools.zip_longest(chars, chars[1:]):
            if page is not None:
                char["page"] = page
            collected.append(char)

            if next_char is None:
                break

            if char["text"] in "。" or (char["text"] == "." and not next_char["text"].isdigit()):
                full_text_list.append(collected)
                total_chars_count += len(collected)
                collected = []
        if collected:
            full_text_list.append(collected)
            total_chars_count += len(collected)
        return full_text_list, total_chars_count

    @staticmethod
    def get_full_text_list(sorted_all_elements):
        full_text_list = []
        total_chars_count = 0

        for each in sorted_all_elements:
            if each["element_type"] == "paragraphs":
                _full_text_list, _total_chars_count = ExportHTML.split_chars_on_sentence(
                    each.get("chars", []), page=each["page"]
                )
                full_text_list.extend(_full_text_list)
                total_chars_count += _total_chars_count
            elif each["element_type"] == "tables":
                _unit = each.get("unit")
                if _unit:
                    full_text_list.append([{"text": _unit}])
                for _cells in each.get("cells", {}).values():
                    chars = _cells.get("chars", [])
                    for char in chars:
                        char["page"] = each["page"]
                    full_text_list.append(chars)
                    total_chars_count += len(chars)

        return full_text_list, total_chars_count


def fill_last_row_height(table):
    last_row = table["grid"]["rows"][-1] if table["grid"]["rows"] else 0
    table["last_row_height"] = table["outline"][3] - table["outline"][1] - last_row


def cell_merged(table):
    ret = {}
    for merged in table.get("merged", []):
        for cell in merged[1:]:
            ret[f"{cell[0]}_{cell[1]}"] = merged[0]
    return ret


def check_merged_headers(target_table, result_table):
    row_should_deleted = []
    for each_row in range(len(target_table["grid"]["rows"]) + 1):
        for each_column in range(len(target_table["grid"]["columns"]) + 1):
            if (
                f"{str(each_row)}_{str(each_column)}" in target_table["cells"]
                and f"{str(each_row)}_{str(each_column)}" in result_table["cells"]
            ):
                target_table_cell = target_table["cells"][f"{str(each_row)}_{str(each_column)}"]
                result_table_cell = result_table["cells"][f"{str(each_row)}_{str(each_column)}"]
                if (target_table_cell.get("value") or target_table_cell.get("text")) != (
                    result_table_cell.get("value") or result_table_cell.get("text")
                ):
                    return row_should_deleted
            elif (
                f"{str(each_row)}_{str(each_column)}" not in target_table["cells"]
                and f"{str(each_row)}_{str(each_column)}" not in result_table["cells"]
            ):
                continue
            else:
                return row_should_deleted
        row_should_deleted.append(each_row)
    if len(row_should_deleted) == len(target_table["grid"]["rows"]) + 1:
        return []
    return row_should_deleted


def remove_duplicate_headers(each_table, row_should_deleted):
    cells_removed = []
    source_cells = copy.deepcopy(each_table["cells"])
    for each_remove_row in row_should_deleted:
        for remove_cell in each_table["cells"]:
            if remove_cell.startswith(f"{str(each_remove_row)}_"):
                del source_cells[remove_cell]
                cells_removed.append([int(remove_cell.split("_")[0]), int(remove_cell.split("_")[1])])
    each_table["cells"] = source_cells

    for each_removed_cell in cells_removed:
        for each_merged_cell in each_table["merged"]:
            if each_removed_cell in each_merged_cell:
                each_table["merged"].remove(each_merged_cell)

    return each_table


def check_continuous_merge(result_table, target_table):
    """
    check and adjust target columns with 5px
    """
    target_columns_fixed = []
    for target_idx, each_target_column in enumerate(target_table["grid"]["columns"]):
        for each_result_column in result_table["grid"]["columns"][target_idx:]:
            if abs(each_target_column - each_result_column) <= 5:
                target_columns_fixed.append(each_result_column)
                break
        else:
            target_columns_fixed.append(each_target_column)

    return target_columns_fixed


def get_columns_interval(target_table):
    """
    list of target columns interval
    """
    columns_interval = []
    target_table_columns = copy.deepcopy(target_table["grid"]["columns"])

    target_table_columns.insert(0, 0)
    target_table_columns.append(target_table["outline"][2])

    for idx in range(len(target_table_columns) - 1):
        columns_interval.append(Interval(target_table_columns[idx], target_table_columns[idx + 1], lower_closed=False))

    return columns_interval


def get_cells_map(result_table, target_table, columns_interval, new_row_at=-1):
    result_table_columns = copy.deepcopy(result_table["grid"]["columns"])
    result_table_columns.append(target_table["outline"][2])

    merged_box_position = defaultdict(list)
    for column_idx, each_column in enumerate(result_table_columns):
        for table_idx, each_interval in enumerate(columns_interval):
            if each_column in each_interval:
                _cell_pos = f"{str(new_row_at + 1)}_{str(table_idx)}"
                _sub_cell_pos = f"{str(new_row_at + 1)}_{str(column_idx)}"
                merged_box_position[_cell_pos].append(_sub_cell_pos)

    return merged_box_position


def get_cells_position_and_merge(merged_box_position, current_cells):
    merged_cells_data = []
    new_position_current_cells = {}
    for each_cell, merge_cells in merged_box_position.items():
        if each_cell in current_cells:
            new_position_current_cells[merge_cells[0]] = current_cells[each_cell]
            merge_cells_box = []
            if len(merge_cells) > 1:
                for each_merged in merge_cells:
                    each_merged = each_merged.split("_")
                    merge_cells_box.append([int(each_merged[0]), int(each_merged[1])])
                merged_cells_data.append(merge_cells_box)

    return new_position_current_cells, merged_cells_data


def get_merge_data_details(result_table, target_table, current_cells, new_row_at):
    target_table["grid"]["columns"] = check_continuous_merge(result_table, target_table)
    columns_interval = get_columns_interval(target_table)
    merged_box_position = get_cells_map(result_table, target_table, columns_interval, new_row_at=new_row_at)
    return get_cells_position_and_merge(merged_box_position, current_cells)


def combine_multi_rows2single(result_tbl, target_tbl):
    result_rows = result_tbl["grid"]["rows"]
    target_rows = target_tbl["grid"]["rows"]
    last_row = result_rows[-1] if result_rows else 0
    result_rows.append(last_row + result_tbl["last_row_height"])
    last_row_position = result_rows[-1]
    for each_row in target_rows:
        result_rows.append(last_row_position + each_row)
    result_tbl["last_row_height"] = target_tbl["last_row_height"]


def merge_multi_tables2single(new_row_at, each_table, result_table, row_should_deleted, merge_cells, merge_map):
    current_uuid = each_table.get("uuid")
    merge_row_count = 0

    current_cells = {}
    for cell_pos, cell_data in each_table["cells"].items():
        if current_uuid in merge_cells:
            merge_row_count = 1

            if cell_pos.startswith(f"{str(len(row_should_deleted))}_"):
                _position = f"{str(new_row_at)}_{cell_pos.split('_')[-1]}"
                if _position in merge_map:
                    old_cell = merge_map[_position]
                    _position = f"{old_cell[0]}_{old_cell[1]}"
                if _position not in result_table["cells"]:
                    continue
                result_table["cells"][_position]["value"] += cell_data.get("value", "")
                result_table["cells"][_position]["related_cell"] = {
                    "uuid": cell_data.get("uuid", ""),
                    "origin_cell": cell_data.get("origin_cell", ""),
                }
                continue
        new_row_position = new_row_at + int(cell_pos.split("_")[0]) + 1 - len(row_should_deleted) - merge_row_count
        new_position = f"{str(new_row_position)}_{str(cell_pos.split('_')[-1])}"
        current_cells[new_position] = cell_data

    if merge_row_count:
        each_table_rows_remain = len(each_table["grid"]["rows"]) - merge_row_count
        each_table["grid"]["rows"] = each_table["grid"]["rows"][:each_table_rows_remain]

    if len(result_table["grid"]["columns"]) != len(each_table["grid"]["columns"]) and not each_table["grid"]["rows"]:
        new_position_current_cells, new_merge_data = get_merge_data_details(
            result_table, each_table, current_cells, new_row_at
        )
        result_table["cells"].update(new_position_current_cells)
        result_table["merged"] += new_merge_data
    else:
        new_merge_data = []
        for each_new_merge in each_table["merged"]:
            _merge = []
            for each_pos in each_new_merge:
                if merge_row_count and each_pos[0] == 0:
                    continue

                each_pos[0] += new_row_at + 1 - len(row_should_deleted) - merge_row_count
                _merge.append([each_pos[0], each_pos[1]])
            if _merge:
                new_merge_data.append(_merge)
        result_table["merged"] += new_merge_data
        result_table["cells"].update(current_cells)

    combine_multi_rows2single(result_table, each_table)

    return result_table


def get_result_table_data(table_details_list, merge_cells):
    result_table = table_details_list[0]
    fill_last_row_height(result_table)

    rest_tables = table_details_list[1:]
    for each_table in rest_tables:
        fill_last_row_height(each_table)
        merge_map = cell_merged(result_table)
        rows_should_deleted = {"header": [], "merge_cells": []}
        if len(each_table["grid"]["columns"]) == len(result_table["grid"]["columns"]):
            # check merge headers
            rows_should_deleted["header"] = check_merged_headers(each_table, result_table)
            if rows_should_deleted["header"]:
                each_table = remove_duplicate_headers(each_table, rows_should_deleted["header"])
                each_table_rows_remain = len(each_table["grid"]["rows"]) - len(rows_should_deleted["header"])
                each_table["grid"]["rows"] = each_table["grid"]["rows"][:each_table_rows_remain]

            if result_table.get("continued_cell"):
                max_row_at = max(int(each.split("_")[0]) for each in result_table["cells"])
                _current_first_row = min(int(each.split("_")[0]) for each in each_table.get("cells", {}))
                for _key, each_main_cell in result_table.get("cells", {}).items():
                    if _key.startswith(str(max_row_at)):
                        each_main_cell["text"] += (
                            each_table.get("cells", {})
                            .get("_".join([str(_current_first_row)] + _key.split("_")[1:]), {})
                            .get("text", "")
                        )

                rows_should_deleted["merge_cells"].append(_current_first_row)
                each_table = remove_duplicate_headers(each_table, rows_should_deleted["merge_cells"])
                each_table_rows_remain = len(each_table["grid"]["rows"]) - 1
                each_table["grid"]["rows"] = each_table["grid"]["rows"][:each_table_rows_remain]

        # merge multi tables into single one
        new_row_at = max(int(each.split("_")[0]) for each in result_table["cells"])
        result_table = merge_multi_tables2single(
            new_row_at,
            each_table,
            result_table,
            rows_should_deleted["header"] + rows_should_deleted["merge_cells"],
            merge_cells,
            merge_map,
        )
        if rows_should_deleted["header"]:
            if "merged_headers" not in result_table:
                result_table["merged_headers"] = {}
            result_table["merged_headers"][f"{str(each_table.get('page'))}-{str(each_table['outline'])}"] = {
                "header_at": new_row_at + 1,
                "cross_lines": len(rows_should_deleted["header"]),
            }
        result_table["continued_cell"] = each_table.get("continued_cell", False)

    del result_table["last_row_height"]
    return result_table


class MergeContinuousTable:
    @staticmethod
    def check_columns_length_equal(table_list):
        all_columns_list = [each_table.get("grid", {}).get("columns", []) for each_table in table_list]
        for columns in all_columns_list[1:]:
            if len(columns) != len(all_columns_list[0]):
                return False

        return True

    @staticmethod
    def get_target_columns(table_list):
        MergeContinuousTable.get_accurate_columns_position(table_list)
        all_columns_list = [
            (each_table.get("grid", {}).get("columns", []), idx) for idx, each_table in enumerate(table_list)
        ]
        all_columns_list = sorted(all_columns_list, key=lambda x: len(x[0]))

        target_columns = copy.deepcopy(all_columns_list[-1])
        target_columns[0].append(table_list[target_columns[1]]["outline"][2])

        return target_columns[0]

    @staticmethod
    def get_accurate_columns_position(table_list):
        for each_table in table_list:
            _column_offset = each_table["outline"][0]
            _accurate_column_position = []
            for each_column in each_table["grid"]["columns"]:
                each_column += _column_offset
                _accurate_column_position.append(each_column)
            each_table["grid"]["columns"] = _accurate_column_position

    @staticmethod
    def get_columns_interval(target_table):
        columns_interval = []
        target_table_columns = copy.deepcopy(target_table["grid"]["columns"])

        target_table_columns.insert(0, target_table["outline"][0])
        target_table_columns.append(target_table["outline"][2])

        for idx in range(len(target_table_columns) - 1):
            columns_interval.append(
                Interval(target_table_columns[idx], target_table_columns[idx + 1], lower_closed=False)
            )

        return columns_interval

    @staticmethod
    def calculate_new_merge(target_table, target_columns):
        new_merged_data = []
        current_columns_interval = MergeContinuousTable.get_columns_interval(target_table)
        table_origin_merged = target_table.get("merged", [])
        cells_convert_relationship = {}

        for each_cell in target_table.get("cells", {}).keys():
            cell_info = each_cell.split("_")
            _row = int(cell_info[0])
            _column = int(cell_info[1])

            for each_origin_merge in table_origin_merged:
                if [_row, _column] in each_origin_merge:
                    current_cell_merge = []
                    for each_origin_merge_cell in each_origin_merge:
                        origin_row = each_origin_merge_cell[0]
                        origin_column = each_origin_merge_cell[1]
                        current_cell_merge += MergeContinuousTable.get_new_cell_merge(
                            current_columns_interval, target_columns, origin_column, origin_row
                        )
                    if len(current_cell_merge) > 1:
                        new_merged_data.append(current_cell_merge)
                    if current_cell_merge:
                        cells_convert_relationship[each_cell] = (
                            f"{str(current_cell_merge[0][0])}_{str(current_cell_merge[0][1])}"
                        )
                    break
            else:
                current_cell_merge = MergeContinuousTable.get_new_cell_merge(
                    current_columns_interval, target_columns, _column, _row
                )
                if len(current_cell_merge) > 1:
                    new_merged_data.append(current_cell_merge)
                if current_cell_merge:
                    cells_convert_relationship[each_cell] = (
                        f"{str(current_cell_merge[0][0])}_{str(current_cell_merge[0][1])}"
                    )
                elif new_merged_data:
                    cells_convert_relationship[each_cell] = (
                        f"{str(new_merged_data[-1][-1][0])}_{str(new_merged_data[-1][-1][1] + 1)}"
                    )

        target_table["merged"] = new_merged_data
        if target_table["grid"]["columns"] and target_table["grid"]["columns"][-1] != target_columns[-2]:
            MergeContinuousTable.calc_tail_merge(target_columns, target_table)

        return cells_convert_relationship

    @staticmethod
    def get_new_cell_merge(current_columns_interval, target_columns, column, row):
        merge_cell_index = []
        cell_columns_interval = current_columns_interval[column]
        for idx, each_column_line in enumerate(target_columns):
            if each_column_line in cell_columns_interval:
                merge_cell_index.append(idx)

        current_cell_merge = []
        for each_cell_idx in merge_cell_index:
            current_cell_merge.append([row, each_cell_idx])

        return current_cell_merge

    @staticmethod
    def check_table_single_row(table_details_list):
        for table in table_details_list[1:]:
            if not table["grid"]["rows"]:
                return True

        return False

    @staticmethod
    def adjust_internal_columns(current_columns, target_columns):
        new_columns = []
        for each_current_column in current_columns:
            for each_target_column in target_columns:
                if abs(each_current_column - each_target_column) < 2:
                    new_columns.append(each_target_column)
                    break
            else:
                new_columns.append(each_current_column)

        return new_columns

    @staticmethod
    def calc_tail_merge(target_columns, target_table):
        tail_merge = []
        columns_diff = len(target_columns) - len(target_table["grid"]["columns"]) - 1
        if columns_diff > 0:
            for each_row in range(len(target_table["grid"]["rows"]) + 1):
                _max_column = max(
                    int(each.split("_")[-1]) for each in target_table["cells"] if each.startswith(f"{each_row}")
                )
                _row_merge = [[each_row, _max_column]]
                for _ in range(columns_diff):
                    _row_merge.append([each_row, _row_merge[-1][-1] + 1])
                tail_merge.append(_row_merge)
        target_table["merged"] += tail_merge

    @staticmethod
    def check_continuous_by_columns(table_details_list):
        if MergeContinuousTable.check_table_single_row(table_details_list):
            return

        if not MergeContinuousTable.check_columns_length_equal(table_details_list):
            target_columns = MergeContinuousTable.get_target_columns(table_details_list)
            for each_table in table_details_list:
                if each_table.get("grid", {}).get("columns", []) != target_columns[:-1]:
                    each_table["grid"]["columns"] = MergeContinuousTable.adjust_internal_columns(
                        each_table["grid"]["columns"], target_columns
                    )
                    if False in [bool(each in target_columns) for each in each_table["grid"]["columns"]]:
                        MergeContinuousTable.calc_tail_merge(target_columns, each_table)
                        each_table["grid"]["columns"] = target_columns[:-1]
                        cells_convert_relationship = {}
                    else:
                        cells_convert_relationship = MergeContinuousTable.calculate_new_merge(
                            each_table, target_columns
                        )
                        each_table["grid"]["columns"] = target_columns[:-1]

                    result_table_cells = copy.deepcopy(each_table.get("cells", {}))

                    for _old_cell, _new_cell in cells_convert_relationship.items():
                        if _old_cell == _new_cell:
                            continue

                        current_new_cell = each_table.get("cells", {}).get(_old_cell, {})
                        current_new_cell.update({"new": True})
                        result_table_cells[_new_cell] = current_new_cell
                        if result_table_cells[_old_cell].get("new") is not True:
                            del result_table_cells[_old_cell]

                    each_table["cells"] = result_table_cells


class BuildHTML:
    P_BLANK = re.compile(r"\s{2,}(?=</?(p|tr|td|body|html|div|head|meta|style|span))")

    @staticmethod
    def get_html(all_block_list, all_file=False, preview=False):
        if preview:
            target_template = "preview_pdf.html"
        elif all_file:
            target_template = "export_file.html"
        else:
            target_template = "new_page.html"
        template_path = os.path.join(project_root, "template", target_template)
        page_html = Template(filename=template_path, input_encoding="utf-8").render(all_block_list=all_block_list)
        page_html = BuildHTML.P_BLANK.sub("", page_html)
        return page_html

    @staticmethod
    def generate_catalogue_html(catalogue, translate=False):
        if isinstance(catalogue, dict) and catalogue.get("index") == -1:
            return f"""<ul class={"translate" if translate else "original"}>{
                BuildHTML.generate_catalogue_html(catalogue.get("children", []), translate)
            }</ul>"""

        if catalogue:
            result = ""
            for item in catalogue:
                anchor = "-".join([str(item["page"]), str(item["element"])])
                collapse_arrow = '<span class="collapse-arrow"></span>' if item.get("children", []) else "<span></span>"
                result += f"""<li data-level="{item["level"]}">
                                  {collapse_arrow}
                                  <a href="#{anchor}">{item.get("translate_text", "") if translate else item.get("title", "")}</a>
                              </li>"""
                if item.get("children", []):
                    result += f"""<ul>{BuildHTML.generate_catalogue_html(item["children"], translate)}</ul>"""
            return result
        return ""

    @staticmethod
    def fill_paragraph_anchor(page_elements_list):
        for each_element in page_elements_list:
            each_element["anchor_id"] = "-".join([str(each_element["page"]), str(each_element["index"])])

    @staticmethod
    def merge_elements_crossed_column(page_details, merge_table: bool):
        index_will_remove = []
        for idx, each_element in enumerate(page_details):
            if each_element.get("continued", False):
                if each_element["element_type"] == "paragraphs":
                    target_index = BuildHTML.find_continuous_paragraph(idx, page_details)
                    if target_index is None:
                        continue
                    page_details[idx]["text"] = each_element.get("page_merged_paragraph", {}).get(
                        "text", ""
                    ) or page_details[target_index].get("text", "")
                    page_details[idx]["chars"] += page_details[target_index]["chars"]
                    page_details[idx]["anchor_id"] = "&".join(
                        [page_details[idx].get("anchor_id", ""), page_details[target_index].get("anchor_id", "")]
                    )
                    index_will_remove.append(target_index)
                elif each_element["element_type"] == "tables" and merge_table:
                    if idx in index_will_remove:
                        continue
                    target_index = BuildHTML.get_continuous_tables(idx, page_details)
                    index_will_remove += target_index
                    continuous_table_list = [each_element] + [page_details[each] for each in target_index]
                    MergeContinuousTable.check_continuous_by_columns(continuous_table_list)
                    continuous_anchor_id = "&".join(each["anchor_id"] for each in continuous_table_list)
                    page_details[idx] = get_result_table_data(continuous_table_list, [])
                    page_details[idx]["anchor_id"] = continuous_anchor_id

        result_full_details = []
        for idx, each in enumerate(page_details):
            if idx not in index_will_remove:
                result_full_details.append(each)

        return result_full_details

    @staticmethod
    def find_continuous_paragraph(current_index, page_details):
        while True:
            if current_index + 1 >= len(page_details):
                return None
            if page_details[current_index + 1]["element_type"] == "paragraphs":
                return current_index + 1
            current_index += 1

    @staticmethod
    def get_continuous_tables(current_index, page_details):
        target_index = []
        while True:
            if current_index + 1 >= len(page_details):
                break
            if page_details[current_index + 1]["element_type"] == "tables":
                target_index.append(current_index + 1)
                if not page_details[current_index + 1].get("continued", False):
                    break
            current_index += 1

        return target_index


def overlap_percent(box1, box2):
    return PdfinsightReader.overlap_percent(box1, box2, base="min")


def mid_point_y(box):
    return (box[1] + box[3]) / 2


def get_syllabuses_level(syllabuses):
    index_level_map = {}
    for each in syllabuses:
        if each["level"] >= 10:
            index_level_map[each["element"]] = 9
        else:
            index_level_map[each["element"]] = each["level"]

    return index_level_map


def fill_paragraphs_level(full_elements, index_level_map):
    for each in full_elements.get("paragraphs", []):
        if each.get("index") in index_level_map:
            # TODO: 什么情况下才会有 style 字段?
            if "style" in each:
                each["styles"]["level"] = index_level_map[each["index"]]


def mark_column_top(column, each_column_ele):
    if column[0] == 0:  # 标记每栏最顶部的元素
        min_outline_top = (None, None)  # (outline[1], idx)
        for idx, ele in enumerate(each_column_ele):
            if min_outline_top[0] is None or ele["outline"][1] < min_outline_top[0]:
                min_outline_top = (ele["outline"][1], idx)
        each_column_ele[min_outline_top[1]]["column_top"] = True

    return each_column_ele


def count_header_footer_image(page_elements, image_struct, header_footer_outline, header=False, footer=False):
    _position = 0 if header and not footer else -1
    if page_elements[_position]["element_type"] == "images":
        page_header_footer_outline = header_footer_outline.get(page_elements[_position]["page"])
        if (
            page_header_footer_outline
            and overlap_percent(page_header_footer_outline, page_elements[_position]["outline"]) > 0.1
        ) or mid_point_y(page_elements[_position]["outline"]) < page_elements[_position]["outline"][3]:
            image_struct.append(page_elements[_position]["outline"])


def get_image_will_remove(image_outline, images_will_remove_outline):
    for each in image_outline:
        if each in images_will_remove_outline:
            continue
        _tmp_overlap_count = 0
        for _outline in image_outline:
            if overlap_percent(each, _outline) >= 0.7:
                _tmp_overlap_count += 1
        if _tmp_overlap_count >= 2:
            images_will_remove_outline.append(each)


def remove_header_image(sorted_elements_by_page, header_outline, footer_outline):
    header_image_outline = []
    footer_image_outline = []
    all_sorted_elements = []
    images_will_remove_outline = []

    for page_elements in sorted_elements_by_page.values():
        if page_elements and len(page_elements) > 1:
            count_header_footer_image(page_elements, header_image_outline, header_outline, header=True)
            count_header_footer_image(page_elements, footer_image_outline, footer_outline, footer=True)
    get_image_will_remove(header_image_outline, images_will_remove_outline)
    get_image_will_remove(footer_image_outline, images_will_remove_outline)

    for page_elements in sorted_elements_by_page.values():
        if page_elements:
            if (
                page_elements[0]["element_type"] == "images"
                and page_elements[0]["outline"] in images_will_remove_outline
            ):
                page_elements.pop(0)
            if (
                page_elements
                and page_elements[-1]["element_type"] == "images"
                and page_elements[-1]["outline"] in images_will_remove_outline
            ):
                page_elements.pop(-1)
            all_sorted_elements += page_elements

    return all_sorted_elements


def remove_tiny_image(all_sorted_elements):
    tiny_images_will_remove_index = []
    for idx, element in enumerate(all_sorted_elements):
        if element:
            if (
                element["element_type"] == "images"
                and element["outline"][2] - element["outline"][0] <= 5
                and element["outline"][3] - element["outline"][1] <= 5
            ):
                if tiny_images_will_remove_index:
                    if tiny_images_will_remove_index[-1][-1] + 1 == idx:
                        tiny_images_will_remove_index[-1].append(idx)
                        continue
                tiny_images_will_remove_index.append([idx])
    tiny_images_will_remove_index = list(filter(lambda x: len(x) > 1, tiny_images_will_remove_index))
    tiny_images_will_remove_index = [idx for _part in tiny_images_will_remove_index for idx in _part]
    filtered_all_elements = []
    for _idx, element in enumerate(all_sorted_elements):
        if _idx not in tiny_images_will_remove_index:
            filtered_all_elements.append(element)

    return filtered_all_elements


def sort_elements(interdoc):
    all_elements_by_pages = {}
    sorted_elements_by_page = OrderedDict()
    include_keys = ["tables", "paragraphs", "images", "shapes", "stamps"]

    header_outline = {each["page"]: each["outline"] for each in interdoc.get("page_headers", [])}
    footer_outline = {each["page"]: each["outline"] for each in interdoc.get("page_footers", [])}

    for each in [header_outline, footer_outline]:
        for _page, _outline in each.items():
            _outline[2] = interdoc.get("pages", {}).get(str(_page), {}).get("size", [595, 841])[0]

    for element, data in interdoc.items():
        if element in include_keys:
            for each_element in data:
                _page = each_element["page"]
                if _page not in all_elements_by_pages:
                    all_elements_by_pages[_page] = []
                each_element["element_type"] = element
                all_elements_by_pages[_page].append(each_element)

    all_elements_by_pages = sorted(all_elements_by_pages.items(), key=lambda x: x[0])

    for page_idx, page_elements in all_elements_by_pages:
        target_elements = {}
        for each in page_elements:
            _position = tuple(each.get("position", [0, 0]) or (0, 0))
            if _position not in target_elements:
                target_elements[_position] = []
            target_elements[_position].append(each)

        for column, each_column_ele in target_elements.items():
            each_column_ele = mark_column_top(column, each_column_ele)
            target_elements[column] = sorted(each_column_ele, key=lambda x: (x.get("outline")[1], x.get("outline")[0]))
        for each in sorted(target_elements.items(), key=lambda x: (x[0][0], x[0][1])):
            if page_idx not in sorted_elements_by_page:
                sorted_elements_by_page[page_idx] = []
            sorted_elements_by_page[page_idx] += each[1]

    for _page, _page_elements in sorted_elements_by_page.items():
        for _element in _page_elements:
            page_size = interdoc.get("pages", {}).get(str(_page), {}).get("size", [])
            if page_size:
                if page_size[0] > page_size[1]:
                    _element["orient"] = PageOrientation.LANDSCAPE.value
                else:
                    _element["orient"] = PageOrientation.PORTRAIT.value
            else:
                _element["orient"] = PageOrientation.PORTRAIT.value

    all_sorted_elements = remove_header_image(sorted_elements_by_page, header_outline, footer_outline)
    return remove_tiny_image(all_sorted_elements)


class BuildWord:
    P_NOT_NUM = re.compile(r"\D+")

    def __init__(self, index_level_map, with_comments):
        self.docx = Document()
        self.index_level_map = index_level_map
        self.with_comments = with_comments

    def add_comment(self, style_block, comment_type):
        # 在指定位置插入批注占位段落；开始段落的前面 和 结束段落的后面
        if not self.with_comments:
            return

        if style_block.get(comment_type):
            for comment_id, comment, user_name in style_block[comment_type]:
                self.docx.add_paragraph(render_comment(comment_id, comment, user_name, comment_type))

    def add_cell_comment(self, style_block, cell):
        # 在cell上插入批注占位段落
        if not self.with_comments:
            return

        if "comment_cell" in style_block:
            for comment_id, comment, user_name in style_block["comment_cell"]:
                cell.add_paragraph(render_comment(comment_id, comment, user_name, "comment_cell"))

    def construct_paragraph(self, block, translate=False):
        style_body = block[0]
        self.add_comment(style_body, "comment_start")
        if style_body.get("level"):
            _header = self.docx.add_paragraph(style=f"Heading {style_body['level']}")
            if style_body.get("text-indent", "0px") != "0px":
                _header.paragraph_format.first_line_indent = Inches(0.3)
            if style_body.get("text-align") == "right":
                _header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            if style_body.get("text-align") == "center":
                _header.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.add_normal_paragraph(block, _header)
            if translate:
                _header = self.docx.add_paragraph(style=f"Heading {style_body['level']}")
                self.add_translate_paragraph(block, _header)
        else:
            self.add_paragraph_with_style(style_body, block, translate=False)
            if translate:
                self.add_paragraph_with_style(style_body, block, translate=True)

        self.add_comment(style_body, "comment_end")

    def add_paragraph_with_style(self, style_body, block, translate=False):
        _paragraph = self.docx.add_paragraph()
        space_before_pt = self.P_NOT_NUM.sub("", style_body.get("margin-top", ""))
        space_before = Pt(int(space_before_pt) if space_before_pt else 5)
        _paragraph.paragraph_format.space_before = space_before
        space_after_pt = self.P_NOT_NUM.sub("", style_body.get("margin-bottom", ""))
        space_after = Pt(int(space_after_pt) if space_after_pt else 5)
        _paragraph.paragraph_format.space_after = space_after
        if style_body.get("text-indent", "0px") != "0px":
            _paragraph.paragraph_format.first_line_indent = Inches(0.3)
        if style_body.get("text-align") == "right":
            _paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if style_body.get("text-align") == "center":
            _paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if not translate:
            self.add_normal_paragraph(block, _paragraph)
        else:
            self.add_translate_paragraph(block, _paragraph)

    @staticmethod
    def add_normal_paragraph(block, target_para):
        for each_span in block[2]:
            span_text = each_span.get("text", "")
            if not span_text:
                continue
            span_text = "".join(x for x in span_text if x.isprintable())
            _run = target_para.add_run(span_text)
            current_span_style = each_span.get("style", {})
            if current_span_style.get("bold"):
                _run.bold = True
            _run.italic = bool(current_span_style.get("italic"))
            _fontsize = current_span_style.get("fontsize", 12)
            _run.size = Pt(_fontsize)
            _fontcolor = hex(current_span_style.get("fontcolor", 0)).strip("0x").rjust(6, "0")
            _run.font.color.rgb = RGBColor(int(_fontcolor[:2], 16), int(_fontcolor[2:4], 16), int(_fontcolor[4:], 16))
            _fontname = current_span_style.get("fontname") or "SimSun"
            _run.font.name = _fontname
            _run_element = _run._element
            _run_element.rPr.rFonts.set(qn("w:eastAsia"), _fontname)

    @staticmethod
    def add_translate_paragraph(block, target_para):
        span_text = "\n".join([each["trans_sen"] for each in block[3]])
        span_text = "".join(x for x in span_text if x.isprintable())
        _run = target_para.add_run(span_text)
        current_span_style = block[2][0].get("style", {}) if block[2] else {}
        if current_span_style:
            if current_span_style.get("bold"):
                _run.bold = True
            _run.italic = bool(current_span_style.get("italic"))
            _fontsize = current_span_style.get("fontsize", 12)
            _run.size = Pt(_fontsize)
            _fontname = "SimSun"
            _fontcolor = hex(current_span_style.get("fontcolor", 0)).strip("0x").rjust(6, "0")
            _run.font.color.rgb = RGBColor(int(_fontcolor[:2], 16), int(_fontcolor[2:4], 16), int(_fontcolor[4:], 16))
            _run.font.name = _fontname
            _run_element = _run._element
            _run_element.rPr.rFonts.set(qn("w:eastAsia"), _fontname)

    def construct_image(self, block):
        _tmp_pic = base64.b64decode(block[1])
        image_file = BytesIO(_tmp_pic)
        self.docx.add_picture(image_file, width=block[0]["width"], height=block[0]["height"])

    @staticmethod
    def process_table_merge(each_cell, current_merge, row, idx):
        if (row, idx) not in current_merge:
            current_merge[(row, idx)] = []
        for row_range in range(row, row + each_cell["rowspan"]):
            for colrange in range(idx, idx + each_cell["colspan"]):
                current_merge[(row, idx)].append((row_range, colrange))

    @staticmethod
    def process_cells_style(_current_style, _cell_run):
        if _current_style.get("fontweight") == "bold":
            _cell_run.bold = True
        if _current_style.get("italic") == "italic":
            _cell_run.italic = True
        _fontname = _current_style.get("fontname") or "SimSun"
        _cell_run.font.name = _fontname
        _run_element = _cell_run._element
        _run_element.rPr.rFonts.set(qn("w:eastAsia"), _fontname)

        return _cell_run

    @staticmethod
    def process_cells_color(_current_style, _cell_run, table, row, idx):
        fontcolor = _current_style.get("fontcolor", "#FFFFFF").strip("#")
        # set text color
        _cell_run.font.color.rgb = RGBColor(int(fontcolor[:2], 16), int(fontcolor[2:4], 16), int(fontcolor[4:], 16))
        bg_color = _current_style.get("bg_color", "#FFFFFF").strip("#")
        shading_elm_1 = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        table.rows[row].cells[idx]._tc.get_or_add_tcPr().append(shading_elm_1)

        return _cell_run, table

    def add_table(self, block, translate=False):
        table_column = max(len(each) for each in block[1])
        table = self.docx.add_table(rows=0, cols=table_column)
        table.style = "Table Grid"
        current_merge = {}

        for row, each in enumerate(block[1]):
            row_cells = table.add_row()
            row_cells.height_rule = WD_ROW_HEIGHT.AUTO
            for idx, each_cell in enumerate(each):
                if "rowspan" in each_cell and "colspan" in each_cell:
                    self.process_table_merge(each_cell, current_merge, row, idx)

        self.set_merge_cells(table, current_merge)
        for row, each in enumerate(block[1]):
            row_cells = table.row_cells(row)
            # row_cells.height_rule = WD_ROW_HEIGHT.AUTO
            for idx, each_cell in enumerate(each):
                _current_style = each_cell.get("styles", {})
                cell_text = (
                    each_cell.get("translate_text", "").strip() if translate else each_cell.get("text", "").strip()
                )
                cell_text = "".join(x for x in cell_text if x.isprintable())
                _cell_run = row_cells[idx].paragraphs[0].add_run(cell_text)

                _cell_run = self.process_cells_style(_current_style, _cell_run)
                _cell_run, table = self.process_cells_color(_current_style, _cell_run, table, row, idx)
                table.cell(row, idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if _current_style.get("align") == "center":
                    row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif _current_style.get("align") == "right":
                    row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                self.add_cell_comment(each_cell, row_cells[idx])

    def construct_table(self, block, translate=False):
        self.add_comment(block[0], "comment_start")
        self.add_table(block)
        if translate:
            self.docx.add_paragraph()
            self.add_table(block, translate=True)
        self.add_comment(block[0], "comment_end")

    def set_merge_cells(self, table, current_merge):
        for target_cell_idx, merge_list in current_merge.items():
            row_count = {each[0] for each in merge_list}
            if len(merge_list) >= 4 and len(row_count) >= 2:
                self.merge_cells_crossed_column(table, merge_list)
            else:
                target_cell = table.cell(target_cell_idx[0], target_cell_idx[1])
                for each_merge_cell in merge_list:
                    target_cell.merge(table.cell(each_merge_cell[0], each_merge_cell[1]))

    @staticmethod
    def merge_cells_crossed_column(table, merge_list):
        row_cells = {}
        for each in merge_list:
            if each[0] not in row_cells:
                row_cells[each[0]] = []
            row_cells[each[0]].append(each)

        for each_cell_list in row_cells.values():
            target_cell = table.cell(each_cell_list[0][0], each_cell_list[0][1])
            for each_cell in each_cell_list:
                target_cell.merge(table.cell(each_cell[0], each_cell[1]))

        result_merge = [each[0] for each in row_cells.values()]
        target_cell = table.cell(result_merge[0][0], result_merge[0][1])
        for each in result_merge:
            target_cell.merge(table.cell(each[0], each[1]))

    def insert_table_of_content_toc(self, translate=False):
        header = self.docx.add_heading("", level=0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header = header.add_run("目录\nCatalogue" if translate else "目录")

        header.bold = True
        header.font.name = "SimSun"
        _run_element = header._element
        _run_element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        # document.add_section(WD_SECTION_START.NEW_PAGE)
        self.create_document_toc()
        # section = document.sections[-1]
        # section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), str(column))

    def create_document_toc(self):
        paragraph = self.docx.add_paragraph()
        run = paragraph.add_run()

        fld_char = OxmlElement("w:fldChar")  # creates a new element
        fld_char.set(qn("w:fldCharType"), "begin")  # sets attribute on element
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")  # sets attribute on element
        instr_text.text = r'TOC \o "1-3" \h \z \u'  # change 1-3 depending on heading levels you need
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        fld_char3 = OxmlElement("w:t")
        fld_char3.text = "Right-click to update field."
        fld_char2.append(fld_char3)
        fld_char4 = OxmlElement("w:fldChar")
        fld_char4.set(qn("w:fldCharType"), "end")
        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)

        self.docx.add_page_break()

    def get_bytes(self):
        if not self.with_comments:
            with BytesIO() as bio:
                self.docx.save(bio)
                return bio.getvalue()

        return replace_comments(self.docx)


def add_new_section(export_file, each_element):
    new_section = export_file.docx.add_section()
    new_section.orientation = WD_ORIENT.LANDSCAPE if each_element[0]["orient"] == 1 else WD_ORIENT.PORTRAIT
    if each_element[0]["orient"] == PageOrientation.PORTRAIT.value:
        new_section.page_width, new_section.page_height = 7772400, 10058400
    else:
        new_section.page_width, new_section.page_height = 10058400, 7772400


def pdf2docx(interdoc: dict, comment_outlines: list = None) -> bytes:
    if comment_outlines:
        # 将 outlines对应到 element index和 cell_index
        comment_positions = get_comment_position_from_interdoc(comment_outlines, interdoc)
        # 在段落和cell上添加待批注标记
        add_comments_in_elements(interdoc, comment_positions)

    index_level_map = get_syllabuses_level(interdoc.get("syllabuses", []))
    fill_paragraphs_level(interdoc, index_level_map)
    sorted_elements = sort_elements(interdoc)
    # TODO:
    # 1. 删除单元格换行符 merge_cell_linebreak
    # 2. 翻译 translate_file
    BuildHTML.fill_paragraph_anchor(sorted_elements)
    sorted_elements = BuildHTML.merge_elements_crossed_column(sorted_elements, True)
    all_block_list = ExportHTML.construct_elements_before_plot(sorted_elements, all_file=True)
    all_block_list = ExportHTML.fix_signed_para_at_end(all_block_list, interdoc["columns"])
    all_block_list = ExportHTML.adjust_char_color(all_block_list)

    export_file = BuildWord(index_level_map, with_comments=bool(comment_outlines))
    # export_file.insert_table_of_content_toc()

    last_element_orient = 0
    for each_element in all_block_list:
        try:
            if last_element_orient != each_element[0]["orient"]:
                last_element_orient = each_element[0]["orient"]
                add_new_section(export_file, each_element)
            if each_element[-1] == "paragraphs":
                export_file.construct_paragraph(each_element)
            elif each_element[-1] == "image":
                export_file.construct_image(each_element)
            elif each_element[-1] == "table":
                export_file.construct_table(each_element)
        except Exception as exp:
            export_file.construct_paragraph(warning_element)
            logger.exception(exp)

    # TODO: 3. 插入目录（需要调用SAAS服务）
    return export_file.get_bytes()
